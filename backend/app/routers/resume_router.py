from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel
from typing import Optional, List, Dict, Any
from pathlib import Path
from uuid import uuid4
import re
import os
import tempfile

router = APIRouter(prefix="/resume", tags=["Resume"])

BASE_DIR = Path(__file__).resolve().parents[2]
UPLOAD_DIR = BASE_DIR / "uploaded_resumes"
GENERATED_DIR = BASE_DIR / "generated_resumes"

UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
GENERATED_DIR.mkdir(parents=True, exist_ok=True)


# -----------------------------
# Helpers
# -----------------------------
def clean_text(text: str) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf_text(file_bytes: bytes) -> str:
    try:
        import fitz
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="PyMuPDF not installed. Run: pip install pymupdf",
        )

    try:
        text = ""
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            for page in doc:
                page_text = page.get_text("text") or ""
                text += page_text + "\n"

        return clean_text(text)

    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"PDF reading failed: {str(e)}",
        )


def extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="python-docx not installed. Run: pip install python-docx",
        )

    tmp_path = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        doc = Document(tmp_path)
        parts = []

        for p in doc.paragraphs:
            if p.text:
                parts.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text
                )
                if row_text:
                    parts.append(row_text)

        return clean_text("\n".join(parts))

    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"DOCX reading failed: {str(e)}",
        )

    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)


def extract_resume_text(file: UploadFile, file_bytes: bytes) -> str:
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return extract_pdf_text(file_bytes)

    if filename.endswith(".docx"):
        return extract_docx_text(file_bytes)

    if filename.endswith(".txt"):
        try:
            return clean_text(file_bytes.decode("utf-8", errors="ignore"))
        except Exception:
            raise HTTPException(status_code=400, detail="TXT reading failed.")

    raise HTTPException(
        status_code=400,
        detail="Unsupported file type. Upload PDF, DOCX, or TXT.",
    )


def normalize_words(text: str) -> List[str]:
    return re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-]*", text.lower())


ROLE_SKILLS: Dict[str, List[str]] = {
    "Frontend Developer": [
        "html", "css", "javascript", "typescript", "react", "next.js",
        "tailwind", "responsive", "api", "github", "ui", "ux",
        "component", "redux", "performance", "accessibility",
    ],
    "Backend Developer": [
        "python", "fastapi", "django", "flask", "node", "express",
        "sql", "postgresql", "mongodb", "api", "rest", "authentication",
        "docker", "redis", "database", "security",
    ],
    "Full Stack Developer": [
        "html", "css", "javascript", "typescript", "react", "next.js",
        "node", "express", "python", "fastapi", "sql", "mongodb",
        "api", "authentication", "github", "deployment",
    ],
    "Data Analyst": [
        "excel", "sql", "python", "statistics", "power bi", "tableau",
        "pandas", "numpy", "visualization", "dashboard", "data cleaning",
        "eda", "business insights",
    ],
    "Data Scientist": [
        "python", "sql", "statistics", "machine learning", "pandas",
        "numpy", "scikit-learn", "model", "regression", "classification",
        "visualization", "feature engineering",
    ],
    "AI/ML Engineer": [
        "python", "machine learning", "deep learning", "tensorflow",
        "pytorch", "nlp", "computer vision", "model", "api",
        "deployment", "mlops", "transformers",
    ],
    "AI Engineer": [
        "python", "machine learning", "deep learning", "tensorflow",
        "pytorch", "nlp", "computer vision", "model", "api",
        "deployment", "mlops", "transformers",
    ],
}


def get_role_skills(role: str) -> List[str]:
    return ROLE_SKILLS.get(role, ROLE_SKILLS["Frontend Developer"])


def analyze_resume_text(text: str, role: str) -> Dict[str, Any]:
    words_text = " ".join(normalize_words(text))
    lower_text = text.lower()
    required = get_role_skills(role)

    found = []
    missing = []

    for skill in required:
        s = skill.lower()
        if s in lower_text or s.replace(".", "") in words_text:
            found.append(skill)
        else:
            missing.append(skill)

    score = int((len(found) / max(len(required), 1)) * 100)

    strengths = []
    improvements = []

    if "project" in lower_text or "projects" in lower_text:
        strengths.append("Projects section detected")
    else:
        improvements.append("Add 2-3 strong projects with measurable outcomes")

    if "github" in lower_text:
        strengths.append("GitHub link or GitHub keyword detected")
    else:
        improvements.append("Add GitHub profile link")

    if "linkedin" in lower_text:
        strengths.append("LinkedIn profile detected")
    else:
        improvements.append("Add LinkedIn profile link")

    if any(x in lower_text for x in ["intern", "experience", "worked", "developed", "built"]):
        strengths.append("Experience/project action words detected")
    else:
        improvements.append("Use action words like built, developed, improved, automated")

    if any(char.isdigit() for char in text):
        strengths.append("Numbers/metrics detected")
    else:
        improvements.append(
            "Add numbers like accuracy, users, time saved, dashboard count, or performance improvement"
        )

    if missing:
        improvements.append(f"Add missing role keywords: {', '.join(missing[:8])}")

    ats_tips = [
        "Use simple headings: Summary, Skills, Projects, Experience, Education, Links.",
        "Avoid tables, heavy graphics, photos, and icons for ATS-friendly resumes.",
        "Add exact keywords from the job description naturally.",
        "Keep fresher resume mostly one page.",
        "Use measurable bullet points: Built X using Y, improved Z by N%.",
    ]

    return {
        "role": role,
        "score": score,
        "level": "Strong Resume"
        if score >= 75
        else "Good Resume"
        if score >= 50
        else "Needs Improvement",
        "found_skills": found,
        "missing_skills": missing,
        "strengths": strengths,
        "improvements": improvements,
        "ats_tips": ats_tips,
        "resume_text": text,
    }


def guess_name(text: str) -> str:
    lines = [x.strip() for x in text.splitlines() if x.strip()]

    for line in lines[:8]:
        if len(line.split()) <= 5 and not any(ch.isdigit() for ch in line):
            if not any(
                x in line.lower()
                for x in ["resume", "curriculum", "email", "phone"]
            ):
                return line.title()

    return ""


def guess_email(text: str) -> str:
    m = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    return m.group(0) if m else ""


def guess_phone(text: str) -> str:
    m = re.search(r"(\+?\d[\d\s\-]{8,}\d)", text)
    return m.group(0).strip() if m else ""


def make_bullets(text: str, fallback: List[str], limit: int = 4) -> List[str]:
    lines = [
        clean_text(x)
        for x in re.split(r"[\n•\-]+", text or "")
        if clean_text(x)
    ]

    final = lines[:limit] if lines else fallback
    return final[:limit]


# -----------------------------
# API Models
# -----------------------------
class AnalyzeTextRequest(BaseModel):
    resume_text: str
    role: Optional[str] = None
    job_role: Optional[str] = None


class GenerateResumeRequest(BaseModel):
    role: Optional[str] = None
    job_role: Optional[str] = None
    template: str = "google"
    resume_format: Optional[str] = "Google ATS"
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    location: Optional[str] = None
    summary: Optional[str] = None
    skills: Optional[str] = None
    projects: Optional[str] = None
    experience: Optional[str] = None
    education: Optional[str] = None
    links: Optional[str] = None
    resume_text: Optional[str] = None


# -----------------------------
# Routes
# -----------------------------
@router.post("/extract")
async def extract_resume(file: UploadFile = File(...)):
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    safe_name = re.sub(
        r"[^a-zA-Z0-9_.-]",
        "_",
        file.filename or f"resume_{uuid4()}",
    )
    saved_path = UPLOAD_DIR / f"{uuid4()}_{safe_name}"
    saved_path.write_bytes(file_bytes)

    text = extract_resume_text(file, file_bytes)

    if len(text) < 20:
        raise HTTPException(
            status_code=400,
            detail="Could not extract enough text. This may be a scanned/image PDF. Use OCR or upload text-based PDF/DOCX.",
        )

    return {
        "success": True,
        "filename": file.filename,
        "saved_path": str(saved_path),
        "text": text,
        "detected": {
            "name": guess_name(text),
            "email": guess_email(text),
            "phone": guess_phone(text),
        },
        "auto_detected": {
            "name": guess_name(text),
            "email": guess_email(text),
            "phone": guess_phone(text),
        },
    }


@router.post("/analyze")
def analyze_resume(payload: AnalyzeTextRequest):
    text = clean_text(payload.resume_text)
    role = payload.role or payload.job_role or "Frontend Developer"

    if len(text) < 20:
        raise HTTPException(status_code=400, detail="Resume text is too short.")

    result = analyze_resume_text(text, role)

    return {
        "job_role": result["role"],
        "resume_format": "Google ATS",
        "score": result["score"],
        "level": result["level"],
        "skills_found": result["found_skills"],
        "missing_skills": result["missing_skills"],
        "strengths": result["strengths"],
        "improvements": result["improvements"],
        "ats_tips": result["ats_tips"],
        "missing_details": [],
    }


@router.post("/analyze-file")
async def analyze_resume_file(
    role: str = Form(...),
    file: UploadFile = File(...),
):
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    text = extract_resume_text(file, file_bytes)

    if len(text) < 20:
        raise HTTPException(
            status_code=400,
            detail="Could not extract enough text. This may be a scanned/image PDF.",
        )

    return analyze_resume_text(text, role)


@router.post("/upload-analyze")
async def upload_analyze_resume(
    file: UploadFile = File(...),
    job_role: str = Form(...),
    resume_format: str = Form("Google ATS"),
):
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="Uploaded file is empty.")

    text = extract_resume_text(file, file_bytes)

    if len(text) < 20:
        raise HTTPException(
            status_code=400,
            detail="Could not extract enough text from this file.",
        )

    result = analyze_resume_text(text, job_role)

    return {
        "filename": file.filename,
        "extracted_text": text,
        "auto_detected": {
            "name": guess_name(text),
            "email": guess_email(text),
            "phone": guess_phone(text),
        },
        "analysis": {
            "job_role": result["role"],
            "resume_format": resume_format,
            "score": result["score"],
            "level": result["level"],
            "skills_found": result["found_skills"],
            "missing_skills": result["missing_skills"],
            "strengths": result["strengths"],
            "improvements": result["improvements"],
            "ats_tips": result["ats_tips"],
            "missing_details": [],
        },
    }


def build_pdf(data: GenerateResumeRequest) -> Path:
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            HRFlowable,
        )
        from reportlab.lib.units import inch
    except Exception:
        raise HTTPException(
            status_code=500,
            detail="reportlab not installed. Run: pip install reportlab",
        )

    raw = data.resume_text or ""

    role = data.role or data.job_role or "Target Role"

    name = data.name or guess_name(raw) or "YOUR NAME"
    email = data.email or guess_email(raw) or "email@example.com"
    phone = data.phone or guess_phone(raw) or "+91 XXXXXXXXXX"
    location = data.location or "City, Country"

    role_skills = get_role_skills(role)
    skills = data.skills or ", ".join(role_skills[:10])

    summary = data.summary or (
        f"Motivated {role} with practical project experience, strong problem-solving ability, "
        f"and role-specific technical skills. Interested in building real-world solutions with clean, reliable execution."
    )

    project_bullets = make_bullets(
        data.projects or "",
        [
            f"Built a practical {role} project using modern tools and clean architecture.",
            "Implemented reusable components, APIs, dashboards, or workflows based on project needs.",
            "Improved usability, performance, and maintainability through structured development.",
        ],
    )

    experience_bullets = make_bullets(
        data.experience or "",
        [
            "Applied technical skills through academic projects, personal projects, and hands-on practice.",
            "Worked with real-world datasets, APIs, UI flows, or backend logic depending on project requirements.",
        ],
        limit=3,
    )

    education = data.education or "Degree / Course Name, Institute Name, Year"
    links = data.links or "LinkedIn | GitHub | Portfolio"

    file_name = f"{role.replace(' ', '_')}_resume_{uuid4()}.pdf"
    output_path = GENERATED_DIR / file_name

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )

    styles = getSampleStyleSheet()

    title = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        alignment=TA_CENTER,
        spaceAfter=6,
    )

    contact = ParagraphStyle(
        "Contact",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        leading=12,
        alignment=TA_CENTER,
        spaceAfter=12,
    )

    heading = ParagraphStyle(
        "Heading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        leading=14,
        alignment=TA_LEFT,
        spaceBefore=8,
        spaceAfter=4,
    )

    body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9.5,
        leading=13,
        alignment=TA_LEFT,
        spaceAfter=4,
    )

    bullet = ParagraphStyle(
        "Bullet",
        parent=body,
        leftIndent=12,
        firstLineIndent=-8,
    )

    story = []

    story.append(Paragraph(name.upper(), title))
    story.append(
        Paragraph(
            f"{email} | {phone} | {location}<br/>{links}",
            contact,
        )
    )
    story.append(HRFlowable(width="100%", thickness=0.8))
    story.append(Spacer(1, 6))

    story.append(Paragraph(role, heading))

    story.append(Paragraph("PROFESSIONAL SUMMARY", heading))
    story.append(Paragraph(summary, body))

    story.append(Paragraph("TECHNICAL SKILLS", heading))
    story.append(Paragraph(skills, body))

    story.append(Paragraph("PROJECTS", heading))
    story.append(Paragraph("Selected Projects", body))

    for b in project_bullets:
        story.append(Paragraph(f"• {b}", bullet))

    story.append(Paragraph("EXPERIENCE", heading))

    for b in experience_bullets:
        story.append(Paragraph(f"• {b}", bullet))

    story.append(Paragraph("EDUCATION", heading))
    story.append(Paragraph(education, body))

    doc.build(story)

    return output_path


@router.post("/generate")
def generate_resume(payload: GenerateResumeRequest):
    pdf_path = build_pdf(payload)

    return {
        "success": True,
        "message": "Resume generated successfully.",
        "file_name": pdf_path.name,
        "filename": pdf_path.name,
        "download_url": f"/resume/download/{pdf_path.name}",
    }


@router.get("/download/{file_name}")
def download_resume(file_name: str):
    safe_name = Path(file_name).name
    file_path = GENERATED_DIR / safe_name

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Resume file not found.")

    return FileResponse(
        path=str(file_path),
        filename=safe_name,
        media_type="application/pdf",
    )